#!/usr/bin/env python3
"""Generate Delta Capita cover letter as a formatted .docx file."""

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
except ImportError:
    print("Installing python-docx...")
    import subprocess
    subprocess.check_call(["pip", "install", "python-docx", "-q"])
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE

def add_paragraph(doc, text, bold=False, font_size=11, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(font_size)
    run.font.name = "Calibri"
    p.paragraph_format.space_after = Pt(space_after)
    return p

def main():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Title
    title = doc.add_paragraph()
    title_run = title.add_run("Application – Senior Frontend Developer (React/TypeScript)")
    title_run.bold = True
    title_run.font.size = Pt(14)
    title_run.font.name = "Calibri"
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(6)

    # Subtitle
    sub = doc.add_paragraph()
    sub_run = sub.add_run("Delta Capita · India (Remote/Hybrid)")
    sub_run.italic = True
    sub_run.font.size = Pt(11)
    sub_run.font.name = "Calibri"
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(18)

    # Candidate name block
    add_paragraph(doc, "Shubham Printe", bold=True, font_size=11, space_after=2)
    add_paragraph(doc, "[Your Phone]", space_after=2)
    add_paragraph(doc, "[Your Email]", space_after=2)
    add_paragraph(doc, "[LinkedIn / Portfolio – optional]", space_after=18)

    # Salutation
    add_paragraph(doc, "Dear Hiring Manager,", space_after=12)

    # Intro
    add_paragraph(
        doc,
        "I am writing to apply for the Senior Frontend Developer position (React & TypeScript, SaaS frontend) at Delta Capita, based in India with remote or hybrid working. I am keen to contribute to building scalable, high-performance frontends and rich UI/UX experiences, and this role aligns well with my experience and career goals.",
        space_after=18,
    )

    # Section: Relevant Experience
    p_heading = doc.add_paragraph()
    run_heading = p_heading.add_run("Relevant Experience")
    run_heading.bold = True
    run_heading.font.size = Pt(12)
    run_heading.font.name = "Calibri"
    p_heading.paragraph_format.space_after = Pt(10)

    bullets = [
        ("React & TypeScript", "[X] years of experience building and maintaining reusable, modular, and high-performance UI components in production applications."),
        ("Mobile Web Development", "Proven experience developing responsive, mobile-first web applications with a focus on performance and usability."),
        ("Performance & Optimization", "Hands-on experience with Vite for fast development and production builds, along with lazy loading, code-splitting, and optimization techniques to ensure smooth rendering."),
        ("API Integration", "Experience connecting frontend applications to RESTful APIs for dynamic data rendering and real-time updates."),
        ("Rich UI/UX", "Designing and implementing interactive, visually appealing interfaces with animations, transitions, and intuitive user flows."),
        ("Testing & Code Quality", "Writing unit and integration tests with Jest and React Testing Library, and maintaining clean, maintainable code."),
    ]
    for label, text in bullets:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f"{label}: ").bold = True
        p.add_run(text)
        for run in p.runs:
            run.font.name = "Calibri"
            run.font.size = Pt(11)
        p.paragraph_format.space_after = Pt(6)

    doc.add_paragraph()  # spacing

    # Closing block
    add_paragraph(
        doc,
        "I am comfortable working in English, am based in India with valid work authorization, and am happy to discuss any reasonable adjustments during the application or interview process if needed.",
        space_after=12,
    )
    add_paragraph(
        doc,
        "I have attached my CV and would welcome the opportunity to discuss how my experience can contribute to Delta Capita's technology and consulting offerings.",
        space_after=12,
    )
    add_paragraph(doc, "Thank you for considering my application.", space_after=18)
    add_paragraph(doc, "Best regards,", space_after=2)
    add_paragraph(doc, "Shubham Printe", bold=True, space_after=0)

    out_path = __file__.replace("generate_cover_letter_docx.py", "Cover-Letter-Delta-Capita.docx")
    doc.save(out_path)
    print(f"Saved: {out_path}")

if __name__ == "__main__":
    main()
