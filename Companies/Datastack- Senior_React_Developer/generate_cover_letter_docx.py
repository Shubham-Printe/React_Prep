#!/usr/bin/env python3
"""Generate Senior React Developer cover letter as a formatted .docx file."""

try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    import subprocess
    subprocess.check_call(["pip", "install", "python-docx", "-q"])
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_paragraph(doc, text, bold=False, font_size=11, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(font_size)
    run.font.name = "Calibri"
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_paragraph_with_bolds(doc, parts, font_size=11, space_after=6):
    """parts: list of (text, bold) tuples"""
    p = doc.add_paragraph()
    for text, bold in parts:
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(font_size)
        run.font.name = "Calibri"
    p.paragraph_format.space_after = Pt(space_after)
    return p


def main():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Title
    title = doc.add_paragraph()
    r = title.add_run("Cover Letter – Senior React Developer / React Team Lead")
    r.bold = True
    r.font.size = Pt(14)
    r.font.name = "Calibri"
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(6)

    # Subtitle
    sub = doc.add_paragraph()
    r = sub.add_run("[Company Name]")
    r.italic = True
    r.font.size = Pt(11)
    r.font.name = "Calibri"
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(18)

    # Contact
    add_paragraph(doc, "Shubham Printe", bold=True, space_after=2)
    add_paragraph(doc, "+91 9011600097", space_after=2)
    add_paragraph(doc, "shubhamprinte99@gmail.com", space_after=18)

    # Body
    add_paragraph(doc, "Dear Hiring Manager,", space_after=12)

    add_paragraph_with_bolds(
        doc,
        [
            ("I am writing to apply for the ", False),
            ("Senior React Developer / React Expert", True),
            (" position. The opportunity to own user-interface development, drive React best practices, and grow into leading the React development team strongly aligns with my experience building production frontends with React.js and my aspiration to provide technical leadership.", False),
        ],
        space_after=12,
    )

    add_paragraph_with_bolds(
        doc,
        [
            ("I have ", False),
            ("[X]+ years", True),
            (" of hands-on experience as a React developer, with a strong focus on ", False),
            ("component architecture", True),
            (", ", False),
            ("state management", True),
            (", and ", False),
            ("code quality", True),
            (". I have designed and implemented reusable component libraries and front-end solutions using ", False),
            ("React", True),
            (", ", False),
            ("Redux", True),
            (", and ", False),
            ("Context API", True),
            (", and I am comfortable with modern workflows including Flux-style data flow. I write ", False),
            ("clean, maintainable code", True),
            (" and regularly ", False),
            ("conduct code reviews", True),
            (" to uphold coding standards and share best practices with the team. I am used to ", False),
            ("integrating front-end applications with back-end services and RESTful APIs", True),
            (", and to debugging complex issues in asynchronous, data-driven UIs.", False),
        ],
        space_after=12,
    )

    add_paragraph_with_bolds(
        doc,
        [
            ("I am proficient in ", False),
            ("JavaScript (ES6+)", True),
            (", ", False),
            ("HTML5", True),
            (", ", False),
            ("CSS3/SASS", True),
            (", and ", False),
            ("responsive design", True),
            (", and I work with ", False),
            ("modern build pipelines and tools", True),
            (" such as Webpack, Babel, and npm/yarn. I have experience ", False),
            ("optimizing existing features and applications", True),
            (" for performance and scalability, and I collaborate closely with other layers of the infrastructure—backend, design, and product—to deliver seamless, high-performance user experiences. I am also experienced with ", False),
            ("version control (Git)", True),
            (" and ", False),
            ("Agile/Scrum", True),
            (" methodologies and value clear communication and teamwork in a collaborative environment.", False),
        ],
        space_after=12,
    )

    add_paragraph_with_bolds(
        doc,
        [
            ("I am eager to ", False),
            ("provide technical leadership and guidance", True),
            (" to the React development team, to ", False),
            ("define and enforce best practices", True),
            (" for front-end development—including coding standards, code reviews, and performance optimization—and to grow toward ", False),
            ("leading the React developer team", True),
            (". I would welcome the opportunity to discuss how my experience in React.js, state management, and technical collaboration can support your product and engineering goals.", False),
        ],
        space_after=12,
    )

    add_paragraph(doc, "Thank you for considering my application.", space_after=18)
    add_paragraph(doc, "Sincerely,", space_after=2)
    add_paragraph(doc, "Shubham Printe", bold=True, space_after=0)

    out_path = __file__.replace("generate_cover_letter_docx.py", "Cover-Letter-Senior-React-Developer.docx")
    doc.save(out_path)
    print(f"Saved: {out_path}")


if __name__ == "__main__":
    main()
