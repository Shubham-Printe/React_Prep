#!/usr/bin/env python3
"""Generate SavvyMoney cover letter as a formatted .docx file."""

try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    import subprocess
    subprocess.check_call(["pip", "install", "python-docx", "-q"])
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_paragraph(doc, text, bold=False, font_size=11, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(font_size)
    run.font.name = "Calibri"
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_paragraph_with_bolds(doc, parts, font_size=11, space_after=6):
    """parts: list of (text, bold) tuples"""
    p = doc.add_paragraph()
    for text, bold in parts:
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(font_size)
        run.font.name = "Calibri"
    p.paragraph_format.space_after = Pt(space_after)
    return p


def main():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Title
    title = doc.add_paragraph()
    r = title.add_run("Cover Letter – Software Development Engineer, React JS / Front End")
    r.bold = True
    r.font.size = Pt(14)
    r.font.name = "Calibri"
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(6)

    # Subtitle
    sub = doc.add_paragraph()
    r = sub.add_run("SavvyMoney · 100% Remote (India)")
    r.italic = True
    r.font.size = Pt(11)
    r.font.name = "Calibri"
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(18)

    # Contact
    add_paragraph(doc, "Shubham Printe", bold=True, space_after=2)
    add_paragraph(doc, "[Your Phone]", space_after=2)
    add_paragraph(doc, "[Your Email]", space_after=18)

    # Body
    add_paragraph(doc, "Dear Hiring Manager,", space_after=12)

    add_paragraph_with_bolds(
        doc,
        [
            ("I am writing to apply for the ", False),
            ("Software Development Engineer, React JS / Front End", True),
            (" position at SavvyMoney. The opportunity to build responsive, user-centric web applications for a leading fintech that serves 1,475+ bank and credit union partners—and to do so as part of a distributed global team from India—strongly aligns with my experience in React/TypeScript and my interest in impactful, scalable products.", False),
        ],
        space_after=12,
    )

    add_paragraph_with_bolds(
        doc,
        [
            ("I have [X]+ years of hands-on experience building responsive web applications with ", False),
            ("React", True),
            (" and ", False),
            ("TypeScript", True),
            (". In my current role, I develop and maintain production frontends with a focus on reusability, performance, and maintainability. I collaborate closely with UX/UI designers to implement designs faithfully, apply UX best practices, and ensure a seamless experience across devices. I am comfortable taking design guidance, iterating on feedback, and translating mockups into clean, accessible interfaces.", False),
        ],
        space_after=12,
    )

    add_paragraph_with_bolds(
        doc,
        [
            ("I am used to ", False),
            ("optimizing applications for speed and scalability", True),
            ("—through code-splitting, lazy loading, and performance profiling—and to ensuring ", False),
            ("cross-browser compatibility and mobile responsiveness", True),
            (". I write clean, well-documented code and take ownership of ", False),
            ("troubleshooting and debugging", True),
            (" to improve both performance and usability. I work regularly with ", False),
            ("backend developers", True),
            (" on API contracts, integration, and end-to-end feature delivery, and I collaborate with product managers and other stakeholders in a multidisciplinary setup. I enjoy breaking down ambiguous problems and driving them to practical solutions, which matches the analytical, problem-solver mindset you described.", False),
        ],
        space_after=12,
    )

    add_paragraph_with_bolds(
        doc,
        [
            ("SavvyMoney's recognition as one of the Top 25 Places to Work in the San Francisco Bay Area and as an Inc. 5000 Fastest Growing Company reflects a culture I want to be part of. I am based in India, set up for ", False),
            ("100% remote work", True),
            (", and would be glad to contribute to your distributed team across the USA, Canada, Europe, and India.", False),
        ],
        space_after=12,
    )

    add_paragraph(
        doc,
        "I would welcome the opportunity to discuss how my experience in React/TypeScript, design collaboration, and performance-focused frontend development can support SavvyMoney's credit score and personal finance solutions.",
        space_after=12,
    )
    add_paragraph(doc, "Thank you for considering my application.", space_after=18)
    add_paragraph(doc, "Sincerely,", space_after=2)
    add_paragraph(doc, "Shubham Printe", bold=True, space_after=0)

    out_path = __file__.replace("generate_cover_letter_docx.py", "Cover-Letter-SavvyMoney.docx")
    doc.save(out_path)
    print(f"Saved: {out_path}")


if __name__ == "__main__":
    main()
